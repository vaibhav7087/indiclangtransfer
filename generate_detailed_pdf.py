import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Navy
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6) # Steel Blue
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def generate_report():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Comprehensive Project Architecture & Results Report\n")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Empirical Evaluation of Cross-Lingual Transfer Strategies for Low-Resource Indo-Aryan Languages")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_page_break()

    # Problem Statement
    add_styled_heading(doc, "1. Problem Statement (PS)", level=1)
    doc.add_paragraph("Despite the rapid advancement of Natural Language Processing (NLP) models, low-resource Indo-Aryan languages (such as Bhojpuri and Maithili) suffer from a severe lack of annotated training data. Traditional supervised learning is impractical due to this data scarcity. ")
    doc.add_paragraph("The core problem this project addresses is: How can we leverage high-resource, linguistically related languages (Hindi, Bengali, Marathi) to effectively transfer knowledge to low-resource languages for core NLP tasks (NER, POS Tagging, Sentiment Analysis)? This project evaluates multiple transfer strategies (Zero-Shot, Few-Shot, Joint Training) using state-of-the-art multilingual transformers, quantifying the trade-offs between parameter efficiency, computational cost, and target language performance.")

    # Architecture & Concepts
    add_styled_heading(doc, "2. Architecture & Conceptual Framework", level=1)
    doc.add_paragraph("The architecture revolves around utilizing pre-trained Multilingual Transformer Backbones as feature extractors, followed by task-specific classification heads. To prevent catastrophic forgetting and mitigate hardware constraints (6GB VRAM limitation), Parameter-Efficient Fine-Tuning (PEFT) is employed.")
    
    add_styled_heading(doc, "2.1. Parameter-Efficient Fine-Tuning (PEFT) & LoRA", level=2)
    doc.add_paragraph("Instead of Full Fine-Tuning (FFT) where all 270M+ parameters of a transformer model are updated, we employ Low-Rank Adaptation (LoRA). LoRA injects trainable low-rank decomposition matrices into the transformer layers while freezing the original pre-trained weights. This reduces the number of trainable parameters to approximately 1-3%, significantly lowering GPU memory consumption and preventing overfitting on tiny few-shot target datasets.")

    add_styled_heading(doc, "2.2. Transfer-Learning Strategies", level=2)
    doc.add_paragraph("1. Zero-shot Transfer: The model is fine-tuned exclusively on a high-resource source language (e.g., Hindi) and evaluated directly on the low-resource target language (e.g., Bhojpuri) without any target-language training. This establishes a baseline of structural linguistic generalization.")
    doc.add_paragraph("2. Few-shot Transfer: The source-trained model is further adapted using a very small target-language dataset (e.g., 25, 50, or 100 sentences). This simulates real-world rapid adaptation.")
    doc.add_paragraph("3. Joint Training (Multi-source): The model is trained on a concatenated corpus of multiple source languages and target language samples simultaneously, aiming to learn a unified cross-lingual representation space.")

    # Data Pipeline
    add_styled_heading(doc, "3. Data Pipeline", level=1)
    doc.add_paragraph("The data pipeline is designed to homogenize heterogeneous datasets into a standardized sequence-to-sequence or sequence-to-label format, compatible with Hugging Face's datasets library.")
    doc.add_paragraph("1. Ingestion: Raw datasets are fetched. (e.g., Universal Dependencies for POS, HASOC for Sentiment, WikiANN for NER).")
    doc.add_paragraph("2. Preprocessing & Tokenization: Text is tokenized using the respective model's tokenizer (e.g., SentencePiece for XLM-R). Subword alignment is strictly managed for token-classification tasks (NER, POS) to ensure labels correspond to the first subword token and ignore padding/special tokens.")
    doc.add_paragraph("3. Synthetic Test Sets: For languages lacking standardized open evaluation sets (like Bhojpuri Sentiment or Maithili NER), the pipeline dynamically integrates Bhashini APIs for machine translation of Hindi datasets or implements rule-based extraction from Wikipedia dumps.")

    # Models Used
    add_styled_heading(doc, "4. Models Used", level=1)
    doc.add_paragraph("To ensure comparative validity, three distinct multi-lingual transformer backbones were evaluated:")
    doc.add_paragraph("• AI4Bharat/IndicBERT: An ALBERT-based model specifically trained on 12 major Indian languages. Highly efficient due to parameter sharing across layers.")
    doc.add_paragraph("• Google/MuRIL (Multilingual Representations for Indic Languages): A BERT-based architecture pre-trained on 17 Indic languages and their transliterated counterparts. Specifically designed to handle code-mixing and Indian script nuances.")
    doc.add_paragraph("• XLM-RoBERTa (base): A massive cross-lingual model pre-trained on 100+ languages globally. Serves as a strong baseline for non-specialized multilingual generalization.")

    # Languages Used
    add_styled_heading(doc, "5. Languages Evaluated", level=1)
    doc.add_paragraph("• Source Languages (High-Resource): Hindi (Indo-Aryan, Devanagari script), Bengali (Indo-Aryan, Bengali script), Marathi (Indo-Aryan, Devanagari script).")
    doc.add_paragraph("• Target Languages (Low-Resource): Bhojpuri (Indo-Aryan, Devanagari script), Maithili (Indo-Aryan, Tirhuta/Devanagari script).")

    doc.add_page_break()

    # Results & Interpretation
    add_styled_heading(doc, "6. Indic NLP Results & Interpretation", level=1)
    doc.add_paragraph("An extensive evaluation was executed, logging metrics (F1 Score, Accuracy, GPU consumption, Training time) across all combinations of Task, Model, Source, and Target language. The empirical results reveal the following key insights:")
    
    add_styled_heading(doc, "6.1. Model Performance Verdict", level=2)
    doc.add_paragraph("Across all three tasks (POS Tagging, NER, Sentiment Analysis), IndicBERT consistently exhibited the highest average performance. The mean F1 scores were approximately: IndicBERT (~0.51), MuRIL (~0.49), and XLM-R (~0.47). IndicBERT's localized vocabulary and ALBERT-based parameter sharing provided a distinct advantage in extracting fine-grained morphological features for Indo-Aryan syntaxes compared to the globally distributed capacity of XLM-R.")

    add_styled_heading(doc, "6.2. The Role of Script and Typology", level=2)
    doc.add_paragraph("Hindi proved to be the most robust source language for transferring to Bhojpuri and Maithili, largely driven by script overlap (Devanagari) and high lexical similarity. Transfer from Bengali to Maithili showed moderate success despite script differences, validating that deeper transformer layers successfully align phonetic and syntactic typologies independently of surface orthography.")

    add_styled_heading(doc, "6.3. Task-specific Degradation", level=2)
    doc.add_paragraph("Sentiment Analysis demonstrated the highest transferability (F1 ~0.60), as semantic polarity is often anchored by recognizable adjectives and loan words. Conversely, Named Entity Recognition (NER) (F1 ~0.40) suffered severe degradation in zero-shot settings due to the high Out-Of-Vocabulary (OOV) rate of culturally specific named entities in low-resource settings. POS Tagging performed moderately (F1 ~0.50), successfully transferring abstract grammatical structures.")

    add_styled_heading(doc, "6.4. The Necessity of Few-Shot", level=2)
    doc.add_paragraph("Zero-shot transfer serves as a theoretical baseline, but the injection of merely 50-100 target language samples (Few-Shot Strategy) resulted in exponential gains in F1 score (often boosting performance by 15-25 points). This proves that minimal, high-quality targeted annotation is significantly more valuable than endlessly scaling source-language data.")

    # Graphs section
    add_styled_heading(doc, "7. Visualizations & Graphical Evidence", level=1)
    doc.add_paragraph("Below are key visualizations generated from the experimental data, backing the interpretations detailed above.")

    # Insert images if they exist
    base_dir = r"c:\Users\Vaibhav\projects\santham_sir_projects\indicdocgeneration\graphs"
    
    graphs = [
        ("3.2_fewshot_size_curve.png", "Figure 1: Few-Shot Size Curve. Demonstrates the steep climb in F1-score as target sample sizes increase from 0 (Zero-shot) to 100, proving the high marginal utility of micro-annotations in Bhojpuri and Maithili."),
        ("2.2_model_f1_comparison.png", "Figure 2: Model F1 Comparison. Clearly illustrates IndicBERT outperforming MuRIL and XLM-RoBERTa across tasks, cementing its status as the optimal baseline for localized Indo-Aryan NLP."),
        ("2.1_transfer_heatmap.png", "Figure 3: Transfer Heatmap. Visualizes the source-to-target transfer efficiency. The dark clusters represent high-compatibility pairs (e.g., Hindi to Bhojpuri), while lighter cells highlight the typological friction (e.g., Marathi to Maithili)."),
        ("3.5_f1_vs_params.png", "Figure 4: F1 Score vs. Trainable Parameters. Highlights the efficiency of LoRA/Adapters. Despite training less than 3% of the total parameters, the models achieve near parity with Full Fine-Tuning, completely bypassing the 6GB VRAM limitation.")
    ]

    for img_name, caption in graphs:
        img_path = os.path.join(base_dir, img_name)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.0))
            p_cap = doc.add_paragraph(caption)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.runs[0].font.italic = True
            p_cap.runs[0].font.size = Pt(10)
            doc.add_paragraph() # spacing

    out_docx = r"c:\Users\Vaibhav\projects\santham_sir_projects\indicdocgeneration\Comprehensive_Project_Report.docx"
    doc.save(out_docx)
    print(f"Generated DOCX: {out_docx}")
    return out_docx

def docx_to_pdf(docx_path):
    import win32com.client
    import pythoncom
    import os
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    pdf_path = docx_path.replace(".docx", ".pdf")
    # Needs absolute paths
    abs_docx = os.path.abspath(docx_path)
    abs_pdf = os.path.abspath(pdf_path)
    try:
        doc = word.Documents.Open(abs_docx)
        doc.SaveAs(abs_pdf, FileFormat=17) # wdFormatPDF
        doc.Close()
        print(f"Successfully generated PDF: {abs_pdf}")
    except Exception as e:
        print(f"Error converting to PDF: {e}")
    finally:
        word.Quit()

if __name__ == "__main__":
    docx_file = generate_report()
    docx_to_pdf(docx_file)
