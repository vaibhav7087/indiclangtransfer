import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.runs[0]
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Navy
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6) # Steel Blue
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def build_prd_docx():
    doc = Document()
    
    # Page setup - 0.75 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("Product Requirements Document (PRD)")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Empirical Evaluation of Cross-Lingual Transfer Strategies for Low-Resource Indo-Aryan Languages (Updated Scope)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    # Metadata Box
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(14)
    run_meta = p_meta.add_run("Focus: Objective 2 Implementation  |  Prepared: August 2026  |  Status: Final Revised Scope")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10)
    run_meta.font.bold = True
    run_meta.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    # 1. Project Context
    add_styled_heading(doc, "1. Project Context", level=1)
    p = doc.add_paragraph(
        "This PRD documents the implementation plan for the research project \"A Comprehensive and Linguistically Informed Evaluation of Multilingual Language Models for Cross-Lingual NLP in Low-Resource Indo-Aryan Languages,\" with primary emphasis on Objective 2. "
        "This updated document reconciles earlier task and dataset assumptions: chunking has been eliminated, and the restricted IIT(BHU) NER dataset has been replaced with open WikiANN NER and Sentiment Analysis. "
        "The evaluation scope spans three source languages (Hindi, Bengali, Marathi), two low-resource target languages (Bhojpuri, Maithili), three core tasks (POS Tagging, Sentiment Analysis, NER via WikiANN), and three pre-trained multilingual transformer backbones (IndicBERT, MuRIL, XLM-R)."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # 2. Study Objectives Alignment
    add_styled_heading(doc, "2. Study Objectives Alignment", level=1)
    
    t_obj = doc.add_table(rows=1, cols=3)
    t_obj.style = 'Table Grid'
    t_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_obj = ["#", "Objective (as stated)", "Status / Relevance to this PRD"]
    hdr_cells = t_obj.rows[0].cells
    for i, h in enumerate(headers_obj):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    obj_data = [
        ("1", "Review existing literature on cross-lingual transfer & multilingual embeddings; evaluate how existing models support knowledge transfer from high- to low-resource Indo-Aryan languages.", "Completed — 35+ papers reviewed. No further action needed for this PRD."),
        ("2", "Perform an empirical evaluation of cross-lingual transfer strategies for core NLP tasks in low-resource Indo-Aryan languages.", "PRIMARY FOCUS of this document. Detailed implementation plan and 4 target evaluation matrices below."),
        ("3", "Develop benchmark datasets, standard evaluation protocols, and visualization tools for comparing NLP system performance across low-resource Indo-Aryan languages.", "Partially supported as a by-product of Objective 2's pipeline (results tables/plots); full benchmark packaging is future work."),
        ("4", "Analyse and suggest how linguistic factors affect cross-lingual performance; identify baseline results and key challenges in processing low-resource Indo-Aryan languages.", "Sequentially dependent on Objective 2's results; execution deferred until Objective 2 evaluation data is fully populated.")
    ]
    for row in obj_data:
        r = t_obj.add_row().cells
        r[0].text, r[1].text, r[2].text = row
        set_cell_background(r[0], "F2F2F2")

    # 3. Requirement Decomposition
    add_styled_heading(doc, "3. Objective 2 — Requirement Decomposition", level=1)
    
    t_req = doc.add_table(rows=1, cols=3)
    t_req.style = 'Table Grid'
    t_req.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_req = ["Requirement in Objective Wording", "What it Demands", "How This Revised Plan Satisfies It"]
    hdr_cells = t_req.rows[0].cells
    for i, h in enumerate(headers_req):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    req_data = [
        ('"Empirical evaluation"', "Real fine-tuning experiments, measured quantitative results — not simulations or proposals only.", "Satisfied — actual fine-tuning runs using IndicBERT, MuRIL, and XLM-R, evaluated on held-out test sets across F1 and Accuracy."),
        ('"Cross-lingual transfer strategies" (plural)', "More than one transfer method, systematically compared.", "Satisfied — Zero-shot transfer, Few-shot transfer (200 target samples), and Joint training strategies evaluated for each model."),
        ('"Core NLP tasks" (plural)', "More than one task type covering sequence labeling and classification.", "Satisfied — 3 distinct tasks: POS Tagging (sequence labeling), Sentiment Analysis (sentence classification), and NER via WikiANN (sequence labeling)."),
        ('"Low-resource Indo-Aryan languages" (plural)', "Multiple source and target language pairs establishing high-to-low transfer dynamics.", "Satisfied — 3 High-Resource Sources (Hindi, Bengali, Marathi) transferring to 2 Low-Resource Targets (Bhojpuri, Maithili).")
    ]
    for row in req_data:
        r = t_req.add_row().cells
        r[0].text, r[1].text, r[2].text = row
        set_cell_background(r[0], "F2F2F2")

    # 4. Reconciled Dataset Feasibility & Workarounds
    add_styled_heading(doc, "4. Reconciled Dataset Feasibility & Workarounds", level=1)
    p_feas = doc.add_paragraph(
        "Earlier scope assumed access to the BMM IIT(BHU) NER dataset, which was later verified to be access-gated and copyrighted. "
        "Chunking was also evaluated and eliminated due to lack of dataset availability for Magahi/Maithili. "
        "The framework now establishes 100% operational feasibility by using open datasets alongside standardized test-set generation techniques where direct low-resource datasets are unhosted."
    )
    p_feas.paragraph_format.space_after = Pt(8)
    
    t_ds = doc.add_table(rows=1, cols=4)
    t_ds.style = 'Table Grid'
    t_ds.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_ds = ["Task", "Languages Supported", "Access / Dataset Source", "Methodology & Workarounds"]
    hdr_cells = t_ds.rows[0].cells
    for i, h in enumerate(headers_ds):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    ds_data = [
        ("POS Tagging", "Sources: Hindi (UD), Bengali, Marathi\nTargets: Bhojpuri (UD_Bhojpuri-BHTB), Maithili", "Openly Accessible (UD & BMM Linguistic Resources)", "Bhojpuri is evaluated using UD_Bhojpuri-BHTB. Maithili uses available academic splits or 200-sentence test-set translation from Hindi."),
        ("Sentiment Analysis", "Sources: Hindi (HASOC), Bengali (SentNoB), Marathi (MahaSent)\nTargets: Maithili (SentiMaithili/8K), Bhojpuri", "Openly Accessible on Hugging Face / Bhashini API", "Maithili has native HF sentiment datasets. Bhojpuri sentiment test set is constructed via machine-translation of 200–500 Hindi sentiment samples using Bhashini API."),
        ("NER (WikiANN)", "Sources: Hindi ('hi'), Bengali ('bn'), Marathi ('mr')\nTargets: Bhojpuri ('bh' config in WikiANN), Maithili", "Openly Accessible (WikiANN / PAN-X on Hugging Face)", "Bhojpuri is natively present in WikiANN ('bh'). Maithili zero-shot test evaluation uses 200 sentences extracted from Maithili Wikipedia dump with hyperlink-derived entity labels.")
    ]
    for row in ds_data:
        r = t_ds.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text = row
        set_cell_background(r[0], "F2F2F2")

    # 5. Evaluated Models & Transfer Strategies
    add_styled_heading(doc, "5. Evaluated Models & Transfer Strategies", level=1)
    
    add_styled_heading(doc, "Multilingual Model Backbones (3)", level=2)
    p_m = doc.add_paragraph()
    p_m.add_run("1. IndicBERT: ").bold = True
    p_m.add_run("Language-specific pre-trained ALBERT model covering 12 major Indic languages.\n")
    p_m.add_run("2. MuRIL: ").bold = True
    p_m.add_run("Multilingual Representations for Indic Languages (BERT-base architecture pre-trained on 17 Indic languages + transliterated text).\n")
    p_m.add_run("3. XLM-RoBERTa (XLM-R): ").bold = True
    p_m.add_run("Massive cross-lingual transformer trained on 100+ languages, serving as a global multilingual baseline.")
    
    add_styled_heading(doc, "Transfer Strategies (3)", level=2)
    t_strat = doc.add_table(rows=1, cols=3)
    t_strat.style = 'Table Grid'
    t_strat.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_st = ["Strategy", "Description", "Purpose in Benchmark"]
    hdr_cells = t_strat.rows[0].cells
    for i, h in enumerate(headers_st):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    st_data = [
        ("Zero-shot transfer", "Fine-tune model on source language (Hindi/Bengali/Marathi) only; evaluate directly on target language test set without any target-language training exposure.", "Establishes lower bound of zero-resource cross-lingual transfer."),
        ("Few-shot transfer", "Fine-tune on source language + small target-language samples (e.g., 200 sentences); measure boost from minimal target data.", "Quantifies data-efficiency and rapid adaptation in low-resource settings."),
        ("Joint training", "Fine-tune simultaneously on source language + all available target-language training data.", "Establishes upper bound of multi-task cross-lingual joint fine-tuning.")
    ]
    for row in st_data:
        r = t_strat.add_row().cells
        r[0].text, r[1].text, r[2].text = row
        set_cell_background(r[0], "F2F2F2")

    # 6. Evaluation Framework & Target Results Matrices
    add_styled_heading(doc, "6. Target Results Matrices (4 Tables)", level=1)
    doc.add_paragraph(
        "The core empirical deliverable for Objective 2 consists of 3 task-specific evaluation sub-tables (18 rows each) and 1 Master Summary Conclusion Table. "
        "Both F1 Score and Accuracy are recorded for every model, language pair, and strategy."
    )
    
    sources = ['Hindi', 'Bengali', 'Marathi']
    targets = ['Bhojpuri', 'Maithili']
    models = ['IndicBERT', 'MuRIL', 'XLM-R']
    
    def create_subtable(doc, title, note=None):
        add_styled_heading(doc, title, level=2)
        if note:
            p_n = doc.add_paragraph()
            p_n.add_run(f"Note: {note}").font.italic = True
            p_n.paragraph_format.space_after = Pt(4)
            
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['Source Lang', 'Target Lang', 'Model', 'Zero-Shot F1', 'Zero-Shot Acc', 'Few-Shot F1', 'Few-Shot Acc', 'Joint F1', 'Joint Acc']
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_background(hdr_cells[i], "2E75B6")
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        for src in sources:
            for tgt in targets:
                for mod in models:
                    row_cells = table.add_row().cells
                    row_cells[0].text = src
                    row_cells[1].text = tgt
                    row_cells[2].text = mod
                    for idx in range(3, 9):
                        p = row_cells[idx].paragraphs[0]
                        p.text = "—"
                        if p.runs:
                            p.runs[0].font.size = Pt(8.5)
                    for cell in row_cells[:3]:
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    create_subtable(doc, "Table 6.1: POS Tagging Performance Sub-Table")
    create_subtable(doc, "Table 6.2: Sentiment Analysis Performance Sub-Table")
    create_subtable(doc, "Table 6.3: NER (WikiANN) Performance Sub-Table", note="Maithili zero-shot test set derived via Wikipedia dump extraction or micro-annotation.")
    
    # Master Conclusion Table
    add_styled_heading(doc, "Table 6.4: Main Conclusion Master Table (Rollup)", level=2)
    p_m_desc = doc.add_paragraph("This table aggregates findings from all three sub-tables, highlighting the best performing source language and peak scores for each model across all three strategies.")
    p_m_desc.paragraph_format.space_after = Pt(4)
    
    main_table = doc.add_table(rows=1, cols=8)
    main_table.style = 'Table Grid'
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_main = ['Task', 'Target Lang', 'Model', 'Best Source Lang', 'Peak Zero-Shot (F1/Acc)', 'Peak Few-Shot (F1/Acc)', 'Peak Joint (F1/Acc)', 'Key Finding / Conclusion']
    hdr_cells = main_table.rows[0].cells
    for i, h in enumerate(headers_main):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    tasks = ['POS Tagging', 'Sentiment Analysis', 'NER (WikiANN)']
    for t in tasks:
        for tgt in targets:
            for i, mod in enumerate(models):
                row_cells = main_table.add_row().cells
                if i == 0:
                    row_cells[0].text = t
                    row_cells[1].text = tgt
                row_cells[2].text = mod
                for idx in range(3, 8):
                    p = row_cells[idx].paragraphs[0]
                    p.text = "—"
                    if p.runs:
                        p.runs[0].font.size = Pt(8.5)
                for cell in row_cells[:3]:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    # 7. Single-Session Workflow & Execution Steps
    add_styled_heading(doc, "7. Implementation Workflow & Execution Steps", level=1)
    
    t_wf = doc.add_table(rows=1, cols=3)
    t_wf.style = 'Table Grid'
    t_wf.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_wf = ["Step", "Action Item", "Target Component"]
    hdr_cells = t_wf.rows[0].cells
    for i, h in enumerate(headers_wf):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    wf_data = [
        ("Step 1", "Load & preprocess source datasets: UD Hindi POS, HASOC Hindi Sentiment, WikiANN Hindi/Bengali/Marathi NER.", "Data Pipeline"),
        ("Step 2", "Load & preprocess target datasets: UD Bhojpuri POS, SentiMaithili, WikiANN Bhojpuri ('bh'). Run Bhashini API test-set translation for Bhojpuri Sentiment and Wikipedia extraction for Maithili NER.", "Data Pipeline"),
        ("Step 3", "Fine-tune IndicBERT, MuRIL, and XLM-R backbones on source language training splits for each task.", "Source Model Fine-Tuning"),
        ("Step 4", "Execute Zero-shot evaluation across all 3 source languages to target language test sets.", "Strategy 1 (Zero-Shot)"),
        ("Step 5", "Execute Few-shot fine-tuning using 200 target-language samples per task; measure F1/Accuracy deltas.", "Strategy 2 (Few-Shot)"),
        ("Step 6", "Execute Joint training on combined source + target splits; evaluate upper-bound performance.", "Strategy 3 (Joint Training)"),
        ("Step 7", "Populate Tables 6.1–6.3, compute Master Table 6.4 rollups, and generate transfer loss / gain charts.", "Deliverable Assembly")
    ]
    for row in wf_data:
        r = t_wf.add_row().cells
        r[0].text, r[1].text, r[2].text = row
        set_cell_background(r[0], "F2F2F2")

    # 8. Tools & Infrastructure
    add_styled_heading(doc, "8. Tools & Infrastructure", level=1)
    p_t = doc.add_paragraph()
    p_t.add_run("• Models: ").bold = True
    p_t.add_run("IndicBERT, MuRIL, XLM-RoBERTa (Hugging Face Transformers)\n")
    p_t.add_run("• Core Libraries: ").bold = True
    p_t.add_run("transformers, datasets, seqeval (for POS/NER F1), scikit-learn (for Sentiment Accuracy/F1), IndicNLP library\n")
    p_t.add_run("• Compute Infrastructure: ").bold = True
    p_t.add_run("Google Colab (T4 GPU) / Kaggle Notebooks (P100 GPU) or local CUDA workstation\n")
    p_t.add_run("• Translation API: ").bold = True
    p_t.add_run("Bhashini API / NLLB-200 / Google Translate for test-set creation\n")
    p_t.add_run("• Agent Tooling: ").bold = True
    p_t.add_run("Google Antigravity / Claude Code for autonomous pipeline orchestration")

    # 9. Summary Verdict
    add_styled_heading(doc, "9. Summary Verdict", level=1)
    p_v = doc.add_paragraph(
        "Objective 2 is fully implementable and operationally sound under this updated scope. "
        "By replacing restricted datasets with open WikiANN NER and Sentiment Analysis, evaluating 3 distinct models across 3 source languages and 2 target languages, "
        "and utilizing test-set translation / Wikipedia extraction for missing target datasets, "
        "this plan satisfies all plural requirements of Objective 2 without any access blockers."
    )
    p_v.paragraph_format.space_after = Pt(12)

    output_path = "c:/Users/Vaibhav/projects/santham_sir_projects/indicdocgeneration/Objective2_PRD_Implementation_Plan_Updated.docx"
    doc.save(output_path)
    print(f"Successfully generated updated docx at: {output_path}")
    return output_path

def convert_docx_to_pdf(docx_path, pdf_path):
    import win32com.client
    import pythoncom
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17) # 17 = wdFormatPDF
    doc.Close()
    word.Quit()
    print(f"Successfully generated PDF at: {pdf_path}")

if __name__ == "__main__":
    docx_file = build_prd_docx()
    pdf_file = "c:/Users/Vaibhav/projects/santham_sir_projects/indicdocgeneration/Objective2_PRD_Implementation_Plan_Updated.pdf"
    convert_docx_to_pdf(docx_file, pdf_file)
